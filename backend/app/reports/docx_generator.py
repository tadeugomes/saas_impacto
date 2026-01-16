"""
Gerador de relatórios em formato DOCX.

Cria documentos Word formatados com dados dos indicadores portuários.
"""

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


class DOCXGenerator:
    """Gerador de documentos DOCX para relatórios de módulos."""

    def __init__(self):
        """Inicializa o gerador."""
        self.doc = Document()

    def set_cell_background(self, cell, color: str):
        """Define a cor de fundo de uma célula da tabela."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def add_header(
        self,
        title: str,
        subtitle: str = "",
        porto: str = "",
        ano: Optional[int] = None,
    ):
        """Adiciona o cabeçalho do relatório."""
        # Título principal
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Subtítulo
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

        # Informações do filtro
        if porto or ano:
            filter_para = self.doc.add_paragraph()
            filter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            filter_text = f"Porto: {porto}"
            if ano:
                filter_text += f" | Ano: {ano}"
            filter_run = filter_para.add_run(filter_text)
            filter_run.font.size = Pt(10)
            filter_run.font.italic = True

        # Data de geração
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        self.doc.add_paragraph()  # Espaçamento

    def add_section(self, title: str, level: int = 2):
        """Adiciona uma seção/título."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(14 if level == 2 else 12)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Adiciona borda abaixo do título
        p = para._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_indicator_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        highlight_header: bool = True,
    ):
        """Adiciona uma tabela de indicadores."""
        if not rows:
            self.doc.add_paragraph("Nenhum dado disponível.")
            return

        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.allow_autofit = False

        # Define largura das colunas
        for col in range(len(headers)):
            table.columns[col].width = Inches(5.0 / len(headers))

        # Cabeçalho
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if highlight_header:
                self.set_cell_background(cell, '4472C4')

        # Dados
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, value in enumerate(row):
                cell = row_cells[j]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.doc.add_paragraph()  # Espaçamento após a tabela

    def add_summary_cards(self, cards: List[Dict[str, Any]]):
        """Adiciona cards de resumo em formato de tabela."""
        if not cards:
            return

        # Cria tabela para cards
        table = self.doc.add_table(rows=1, cols=len(cards))
        table.style = 'Light Grid Accent 1'

        for i, card in enumerate(cards):
            cell = table.rows[0].cells[i]
            cell.text = f"{card.get('label', '')}\n\n{card.get('value', '')}"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.set_cell_background(cell, 'D9E2F3')

        self.doc.add_paragraph()

    def add_chart_placeholder(self, chart_title: str):
        """Adiciona um placeholder para gráficos."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Cria uma caixa simulando gráfico
        run = para.add_run(f"\n[Gráfico: {chart_title}]\n\n")
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Borda ao redor
        p = para._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            elm = OxmlElement(f'w:{side}')
            elm.set(qn('w:val'), 'single')
            elm.set(qn('w:sz'), '4')
            pBdr.append(elm)
        pPr.append(pBdr)

    def add_text(self, text: str, bold: bool = False, italic: bool = False):
        """Adiciona um texto simples."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic

    def add_bullet_list(self, items: List[str]):
        """Adiciona uma lista com marcadores."""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')

    def add_page_break(self):
        """Adiciona uma quebra de página."""
        self.doc.add_page_break()

    def save(self) -> BytesIO:
        """Salva o documento em memória e retorna os bytes."""
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def get_filename(self, module: str, porto: str = "", ano: Optional[int] = None) -> str:
        """Gera nome de arquivo para download."""
        parts = [module.lower().replace(' ', '_')]
        if porto:
            parts.append(porto.lower().replace(' ', '_'))
        if ano:
            parts.append(str(ano))
        parts.append(datetime.now().strftime('%Y%m%d_%H%M%S'))
        return '_'.join(parts) + '.docx'
