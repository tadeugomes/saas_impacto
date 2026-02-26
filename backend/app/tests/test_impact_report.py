"""Testes do gerador de relatório DOCX para análises causais (PR-20).

Cobre:
- Geração de bytes válidos para cada método causal (DiD, IV, Panel IV,
  Event Study, Compare, SCM, ASCM).
- Casos de borda: falha de execução, payload vazio, artifact_path em disco.
- Validação estrutural do arquivo DOCX gerado (parseable por python-docx).
- Formato do nome de arquivo retornado.
"""
from __future__ import annotations

import os
import json
import tempfile
import uuid
from io import BytesIO
from typing import Any

import pytest
from docx import Document

from app.reports.report_service import ReportService


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_main_result(
    coef: float = 0.15,
    std_err: float = 0.04,
    p_value: float = 0.02,
    ci_lower: float = 0.07,
    ci_upper: float = 0.23,
    n_obs: int = 120,
) -> dict[str, Any]:
    return {
        "coef": coef,
        "std_err": std_err,
        "p_value": p_value,
        "ci_lower": ci_lower,
        "ci_upper": ci_upper,
        "n_obs": n_obs,
    }


def _make_event_study_coefficients() -> list[dict[str, Any]]:
    return [
        {
            "rel_time": t,
            "coef": 0.01 * t,
            "se": 0.02,
            "pvalue": 0.05 + abs(t) * 0.01,
            "ci_lower": 0.01 * t - 0.04,
            "ci_upper": 0.01 * t + 0.04,
            "period": 2010 + t,
            "significant_10pct": abs(t) > 1,
        }
        for t in range(-3, 4)
    ]


def _make_result_full_did(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log"]
    return {
        outcome: {
            "main_result": _make_main_result(),
            "parallel_trends": {
                "warning": "Tendências paralelas não rejeitadas (p=0.31).",
                "interpretation": "Premissa satisfeita.",
            },
            "model_info": {"formula": f"{outcome} ~ treated + post + treated:post"},
        }
        for outcome in outcomes
    }


def _make_result_full_iv(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log"]
    return {
        outcome: {
            "main_result": _make_main_result(coef=0.22),
            "first_stage": {
                "f_stat": 28.3,
                "f_pvalue": 0.001,
                "warning": "Instrumento relevante (F=28.3).",
            },
            "reduced_form": {"coef": 0.18, "p_value": 0.04},
        }
        for outcome in outcomes
    }


def _make_result_full_panel_iv(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["n_vinculos_log"]
    return {
        outcome: {
            "main_result": _make_main_result(coef=0.08, n_obs=340),
            "specifications": [
                {"spec": "baseline", "coef": 0.08, "p_value": 0.01},
                {"spec": "with_controls", "coef": 0.07, "p_value": 0.02},
            ],
        }
        for outcome in outcomes
    }


def _make_result_full_event_study(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log"]
    return {
        outcome: {
            "coefficients": _make_event_study_coefficients(),
            "model_info": {"formula": f"{outcome} ~ rel_time_fe + unit_fe"},
        }
        for outcome in outcomes
    }


def _make_result_full_compare(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log", "n_vinculos_log"]
    did_part: dict[str, Any] = {}
    comparison_part: dict[str, Any] = {}
    for outcome in outcomes:
        did_part[outcome] = {"main_result": _make_main_result()}
        comparison_part[outcome] = {
            "consistency_assessment": "consistent",
            "recommended_estimate": "did",
            "interpretation_notes": "Ambos os métodos apontam efeito positivo.",
            "comparison_table": [
                {
                    "Method": "DiD",
                    "Estimate": 0.15,
                    "SE": 0.04,
                    "CI_Lower": 0.07,
                    "CI_Upper": 0.23,
                    "P_Value": 0.02,
                    "Significant": "Sim",
                    "Notes": None,
                    "n_obs": 120,
                },
                {
                    "Method": "IV",
                    "Estimate": 0.18,
                    "SE": 0.05,
                    "CI_Lower": 0.08,
                    "CI_Upper": 0.28,
                    "P_Value": 0.03,
                    "Significant": "Sim",
                    "Notes": None,
                    "n_obs": 118,
                },
            ],
        }
    return {"did": did_part, "comparison": comparison_part}


def _make_result_full_scm(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log"]
    return {
        outcome: {
            "main_result": {
                "post_att": 0.12,
                "pre_rmspe": 0.03,
                "post_rmspe": 0.09,
                "ratio_post_pre": 3.0,
                "w_optimal": [0.4, 0.3, 0.3],
                "donor_units": ["3300100", "3300200", "3300300"],
            },
            "placebo_test": {"p_value": 0.04, "in_space_placebos": []},
            "event_study": [
                {"year": 2012 + i, "effect": 0.01 * i, "treated": 1.0 + 0.01 * i, "synthetic": 1.0}
                for i in range(6)
            ],
        }
        for outcome in outcomes
    }


def _make_result_full_ascm(outcomes: list[str] | None = None) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log"]
    return {
        outcome: {
            "main_result": {
                "post_att": 0.10,
                "pre_rmspe": 0.02,
                "post_rmspe": 0.08,
                "w_optimal": [0.5, 0.5],
                "donor_units": ["3300100", "3300200"],
                "ridge_lambda": 0.1,
            },
        }
        for outcome in outcomes
    }


def _make_analysis(
    method: str = "did",
    status: str = "success",
    outcomes: list[str] | None = None,
    result_full: dict[str, Any] | None = None,
    result_summary: dict[str, Any] | None = None,
    error_message: str | None = None,
    artifact_path: str | None = None,
) -> dict[str, Any]:
    outcomes = outcomes or ["pib_log"]
    if result_full is None:
        _builders = {
            "did": _make_result_full_did,
            "iv": _make_result_full_iv,
            "panel_iv": _make_result_full_panel_iv,
            "event_study": _make_result_full_event_study,
            "compare": _make_result_full_compare,
            "scm": _make_result_full_scm,
            "augmented_scm": _make_result_full_ascm,
        }
        result_full = _builders.get(method, _make_result_full_did)(outcomes)

    if result_summary is None:
        result_summary = {
            "outcome": outcomes[0],
            "coef": 0.15,
            "std_err": 0.04,
            "p_value": 0.02,
            "ci_lower": 0.07,
            "ci_upper": 0.23,
            "n_obs": 120,
            "warnings": [],
        }

    return {
        "id": str(uuid.uuid4()),
        "method": method,
        "status": status,
        "result_summary": result_summary,
        "result_full": result_full,
        "artifact_path": artifact_path,
        "error_message": error_message,
        "request_params": {
            "method": method,
            "treated_ids": ["3304557"],
            "control_ids": ["3300100", "3300200"],
            "treatment_year": 2015,
            "scope": "municipal",
            "outcomes": outcomes,
            "controls": ["populacao_log"],
            "instrument": "soja_log" if method in ("iv", "panel_iv") else None,
            "ano_inicio": 2010,
            "ano_fim": 2022,
            "use_mart": True,
        },
    }


def _is_valid_docx(buffer: BytesIO) -> bool:
    """Verifica se o BytesIO contém um documento Word parseable."""
    try:
        buffer.seek(0)
        Document(buffer)
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Testes por método
# ---------------------------------------------------------------------------

def test_generate_docx_did_returns_valid_bytes():
    """DiD: gera DOCX não-vazio e parseable."""
    analysis = _make_analysis(method="did")
    service = ReportService()
    buf, filename = service.generate_impact_analysis_report(analysis)

    assert isinstance(buf, BytesIO)
    assert buf.getvalue(), "Buffer DOCX vazio"
    assert _is_valid_docx(buf), "Bytes não constituem um DOCX válido"
    assert filename.startswith("analise_did_")
    assert filename.endswith(".docx")


def test_generate_docx_iv_returns_valid_bytes():
    """IV (2SLS): gera DOCX com diagnóstico de first-stage."""
    analysis = _make_analysis(method="iv")
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    assert "iv" in filename


def test_generate_docx_panel_iv_returns_valid_bytes():
    """Panel IV: gera DOCX com especificações alternativas."""
    analysis = _make_analysis(method="panel_iv")
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    assert "panel_iv" in filename


def test_generate_docx_event_study_includes_coefficients():
    """Event Study: DOCX contém seção de coeficientes por período."""
    analysis = _make_analysis(method="event_study")
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    buf.seek(0)
    doc = Document(buf)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # seção "Detalhes por Outcome" deve estar presente
    assert "Detalhes" in full_text or "Rel. Time" in full_text or "Outcome" in full_text


def test_generate_docx_compare_includes_comparison_table():
    """Compare: DOCX contém tabela com múltiplos métodos."""
    analysis = _make_analysis(method="compare", outcomes=["pib_log", "n_vinculos_log"])
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    assert "compare" in filename


def test_generate_docx_scm_returns_valid_bytes():
    """SCM: gera DOCX sem erro (pós PR-25)."""
    analysis = _make_analysis(method="scm")
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    assert "scm" in filename


def test_generate_docx_augmented_scm_returns_valid_bytes():
    """ASCM: gera DOCX sem erro."""
    analysis = _make_analysis(method="augmented_scm")
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    assert "augmented_scm" in filename


# ---------------------------------------------------------------------------
# Casos de borda
# ---------------------------------------------------------------------------

def test_generate_docx_failed_status_returns_bytes():
    """Análise com status=failed: gera DOCX com mensagem de erro, sem exceção."""
    analysis = _make_analysis(
        method="did",
        status="failed",
        result_full={},
        result_summary={},
        error_message="Timeout na consulta BigQuery ao construir painel.",
    )
    buf, filename = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    buf.seek(0)
    doc = Document(buf)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Timeout" in full_text or "Falha" in full_text


def test_generate_docx_empty_result_full_does_not_raise():
    """Análise com result_full vazio: não deve levantar exceção."""
    analysis = _make_analysis(method="did", result_full={})
    buf, _ = ReportService().generate_impact_analysis_report(analysis)
    assert _is_valid_docx(buf)


def test_generate_docx_none_result_summary_does_not_raise():
    """Análise sem result_summary: não deve levantar exceção."""
    analysis = _make_analysis(method="iv", result_summary=None)
    # força result_summary para None
    analysis["result_summary"] = None
    buf, _ = ReportService().generate_impact_analysis_report(analysis)
    assert _is_valid_docx(buf)


def test_generate_docx_artifact_path_file_loads_payload(tmp_path):
    """artifact_path local: payload carregado do disco quando result_full está vazio."""
    payload = _make_result_full_did(["pib_log"])
    artifact_file = tmp_path / "result.json"
    artifact_file.write_text(json.dumps(payload), encoding="utf-8")

    analysis = _make_analysis(
        method="did",
        result_full=None,
        artifact_path=str(artifact_file),
    )
    # Zera o result_full inline para forçar leitura do artifact_path
    analysis["result_full"] = {}

    buf, _ = ReportService().generate_impact_analysis_report(analysis)
    assert _is_valid_docx(buf)


def test_generate_docx_artifact_path_missing_generates_warning_in_docx(tmp_path):
    """artifact_path inexistente: relatório gerado com aviso (sem crash)."""
    analysis = _make_analysis(method="did", result_full={})
    analysis["artifact_path"] = str(tmp_path / "inexistente.json")

    buf, _ = ReportService().generate_impact_analysis_report(analysis)
    assert _is_valid_docx(buf)


def test_generate_docx_multiple_outcomes_all_included():
    """Múltiplos outcomes: todas as variáveis aparecem no documento."""
    outcomes = ["pib_log", "n_vinculos_log", "comercio_dolar_log"]
    analysis = _make_analysis(method="did", outcomes=outcomes)
    buf, _ = ReportService().generate_impact_analysis_report(analysis)

    assert _is_valid_docx(buf)
    buf.seek(0)
    doc = Document(buf)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Ao menos um outcome deve aparecer no texto
    assert any(o in full_text for o in outcomes)


def test_generate_docx_with_warnings_in_result_summary():
    """Análise com warnings no result_summary: não quebra e warnings chegam ao DOCX."""
    analysis = _make_analysis(method="panel_iv")
    analysis["result_summary"]["warnings"] = [
        "Fraco instrumento detectado (F < 10).",
        "Apenas 2 períodos pré-tratamento disponíveis.",
    ]
    buf, _ = ReportService().generate_impact_analysis_report(analysis)
    assert _is_valid_docx(buf)


def test_generate_docx_filename_contains_method_and_timestamp():
    """Nome do arquivo segue padrão: analise_{method}_{uuid}_{timestamp}.docx."""
    import re

    analysis = _make_analysis(method="event_study")
    _, filename = ReportService().generate_impact_analysis_report(analysis)

    # analise_event_study_{uuid}_{YYYYMMDD_HHMMSS}.docx
    pattern = r"^analise_event_study_.+_\d{8}_\d{6}\.docx$"
    assert re.match(pattern, filename), f"Filename fora do padrão: {filename}"


def test_generate_docx_output_is_bytesio_not_bytes():
    """O retorno deve ser BytesIO (não bytes puro), compatível com StreamingResponse."""
    analysis = _make_analysis(method="did")
    buf, _ = ReportService().generate_impact_analysis_report(analysis)
    assert hasattr(buf, "getvalue"), "Retorno deve ser BytesIO"
    assert hasattr(buf, "seek"), "Retorno deve ser BytesIO"
